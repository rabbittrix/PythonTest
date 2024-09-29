import streamlit as st
import pdfplumber
import docx
import xml.etree.ElementTree as ET
import json
import re
import os
import datetime
import csv
from collections import OrderedDict
from langdetect import detect, DetectorFactory
from fpdf import FPDF

DetectorFactory.seed = 0

def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_text_from_txt(file):
    try:
        raw_data = file.read()
        encoding = 'utf-8'
        try:
            text = raw_data.decode(encoding)
        except UnicodeDecodeError:
            encoding = 'latin-1'
            text = raw_data.decode(encoding)
    except Exception as e:
        st.error(f"Erro ao ler o arquivo .txt: {e}")
        text = ""
    return text

def extract_text_from_xml(file):
    tree = ET.parse(file)
    root = tree.getroot()
    text = "\n".join([elem.text for elem in root.iter() if elem.text])
    return text

def detect_language(text):
    try:
        lang = detect(text)
        return lang
    except:
        return "Não detectado"

def extract_between_braces(content):
    matches = re.findall(r'\{(.*?)\}', content)
    return matches

def replace_variables(content, replacements):
    highlighted_content = content
    for key, value in replacements.items():
        highlighted_content = re.sub(
            r'\{' + re.escape(key) + r'\}', 
            f'<span style="color:yellow">{value}</span>',  # Alteração em amarelo
            highlighted_content
        )
    # Adiciona o texto original em vermelho
    highlighted_content = re.sub(r'\{(.*?)\}', r'<span style="color:red">{\1}</span>', highlighted_content)
    return highlighted_content

def save_as_docx(content, file_name):
    doc = docx.Document()
    for line in content.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # Texto preto
    doc.save(file_name)

def save_as_pdf(content, file_name):
    pdf = FPDF()
    pdf.add_page()
    
    # Adiciona a fonte Unicode correta
    pdf.add_font("NotoSans", "", "./font/NotoSans-VariableFont_wdth,wght.ttf", uni=True)
    pdf.set_font("NotoSans", size=12)
    
    # Adiciona o conteúdo ao PDF
    pdf.multi_cell(0, 10, content)
    pdf.output(file_name)

def save_as_csv(content, file_name):
    with open(file_name, 'w', newline='') as file:
        writer = csv.writer(file)
        for line in content.split('\n'):
            writer.writerow([line])
            
def save_as_txt(content, file_name):
    with open(file_name, "w", encoding="utf-8") as file:
        file.write(content)
            
def download_file(content, original_file_name, file_format):
    format_mapping = {
        "Texto (.txt)": "txt",
        "DOCX (.docx)": "docx",
        "PDF (.pdf)": "pdf",
        "CSV (.csv)": "csv"
    }
    
    extension = format_mapping.get(file_format, "").lower()
    if not extension:
        st.error("Unsupported file format.")
        return
    
    base_name = os.path.splitext(original_file_name)[0]
    current_date = datetime.datetime.now().strftime("%Y%m%d")
    file_name = f"{base_name}_{current_date}.{extension}"
    
    # Remove HTML tags para exportação
    content_to_save = re.sub(r'<[^>]+>', '', content)
    
    if extension == "pdf":
        save_as_pdf(content_to_save, file_name)
    elif extension == "docx":
        save_as_docx(content_to_save, file_name)
    elif extension == "txt":
        save_as_txt(content_to_save, file_name)
    elif extension == "csv":
        save_as_csv(content_to_save, file_name)

    with open(file_name, "rb") as file:
        st.download_button(
            label=f"Download Document ({extension.upper()})",
            data=file,
            file_name=file_name,
            mime='application/octet-stream'
        )

def main():
    st.title("Data Extraction from Documents")
    
    st.sidebar.title("Upload Document")
    uploaded_file = st.sidebar.file_uploader("Choose a file", type=["pdf", "docx", "txt", "xml"], key="file_uploader")

    if uploaded_file:
        file_type = uploaded_file.name.split('.')[-1]
        content = ""

        if file_type == 'pdf':
            content = extract_text_from_pdf(uploaded_file)
        elif file_type == 'docx':
            content = extract_text_from_docx(uploaded_file)
        elif file_type == 'txt':
            content = extract_text_from_txt(uploaded_file)
        elif file_type == 'xml':
            content = extract_text_from_xml(uploaded_file)
        else:
            st.error("Unsupported file format.")
            return

        if not content:
            st.error("The file contents are empty.")
            return

        variables = extract_between_braces(content)
        
        # Detecção de idioma
        language = detect_language(content)
        st.sidebar.write(f"Language detected: {language}")
        
        # Inputs para as variáveis únicas encontradas
        replacements = {}
        if variables:
            st.sidebar.subheader("Replace Variables")
            for var in OrderedDict.fromkeys(variables):
                replacements[var] = st.sidebar.text_input(var, key=f"input_{var}")

        if st.sidebar.button("Apply Substitutions", key="apply_substitutions"):
            content_to_display = replace_variables(content, replacements)
            st.session_state['content'] = content_to_display  # Armazena o conteúdo atualizado na sessão
            st.session_state['replacements'] = replacements  # Armazena as substituições na sessão

        # Exibe o conteúdo atualizado com cores para mudanças
        content_to_display = st.session_state.get('content', content)
        st.subheader("Document Content")
        st.markdown(content_to_display, unsafe_allow_html=True)  # Exibe com formatação HTML

        # Salva o conteúdo para exportar, removendo tags HTML
        content_to_save = re.sub(r'<[^>]+>', '', content_to_display)
        
        # Verifica se o conteúdo para JSON não está vazio
        if content_to_save.strip():
            if st.button("Convert to JSON", key="convert_to_json"):
                json_content = json.dumps({"content": content_to_save}, indent=4)
                st.json({"content": content_to_save})

                # Opção para download do JSON
                st.download_button(
                    label="Download JSON",
                    data=json_content,
                    file_name='document_content.json',
                    mime='application/json'
                )
        else:
            st.warning("The content is empty and cannot be converted to JSON.")

        # Opção para download do documento alterado
        st.sidebar.subheader("Export Modified Document")
        format_option = st.sidebar.selectbox(
            "Choose the format to export",
            ["Texto (.txt)", "DOCX (.docx)", "PDF (.pdf)", "CSV (.csv)"],
            key="format_option"
        )

        if st.sidebar.button("Export Document", key="export_document"):
            download_file(content_to_save, uploaded_file.name, format_option)

if __name__ == "__main__":
    main()
