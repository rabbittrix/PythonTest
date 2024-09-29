import streamlit as st
import pdfplumber
import docx
import xml.etree.ElementTree as ET
import json
import re
import os
import datetime
import csv
from collections import OrderedDict
from langdetect import detect
from fpdf import FPDF

# Função para extrair texto de PDFs
def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ''
    return text

# Função para extrair texto de arquivos DOCX
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

# Função para extrair texto de arquivos TXT
def extract_text_from_txt(file):
    try:
        raw_data = file.read()
        encoding = 'utf-8'
        try:
            text = raw_data.decode(encoding)
        except UnicodeDecodeError:
            encoding = 'latin-1'
            text = raw_data.decode(encoding)
    except Exception as e:
        st.error(f"Error reading file .txt: {e}")
        text = ""
    return text

# Função para extrair texto de arquivos XML
def extract_text_from_xml(file):
    tree = ET.parse(file)
    root = tree.getroot()
    text = "\n".join([elem.text for elem in root.iter() if elem.text])
    return text

# Função para detectar o idioma do texto
def detect_language(text):
    try:
        lang = detect(text)
        return lang
    except:
        return "Not detected"

# Função para analisar pares de chave-valor no conteúdo do documento
def parse_key_value_pairs(content):
    lines = content.splitlines()
    key_value_pairs = {}
    for line in lines:
        if ": " in line:
            key, value = line.split(": ", 1)
            key_value_pairs[key.strip()] = value.strip()
    return key_value_pairs

# Função para substituir variáveis no conteúdo do documento e destacar apenas os valores alterados
def replace_variables(content, replacements):
    for key, new_value in replacements.items():
        pattern = f'({key}: )(.+)'  # Captura o "key: " e o valor associado
        replacement = f'\\1<span style="color:yellow">{new_value}</span>'
        content = re.sub(pattern, replacement, content)
    return content

# Funções para salvar em diferentes formatos
def save_as_docx(content, file_name):
    doc = docx.Document()
    for line in content.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  
    doc.save(file_name)

def save_as_pdf(content, file_name):
    pdf = FPDF()
    pdf.add_page()
    
    pdf.add_font("NotoSans", "", "./font/NotoSans-VariableFont_wdth,wght.ttf", uni=True)
    pdf.set_font("NotoSans", size=12)
    
    pdf.multi_cell(0, 10, content)
    pdf.output(file_name)

def save_as_csv(content, file_name):
    with open(file_name, 'w', newline='') as file:
        writer = csv.writer(file)
        for line in content.split('\n'):
            writer.writerow([line])

def save_as_txt(content, file_name):
    with open(file_name, "w", encoding="utf-8") as file:
        file.write(content)

# Função para download do arquivo modificado
def download_file(content, original_file_name, file_format):
    format_mapping = {
        "Texto (.txt)": "txt",
        "DOCX (.docx)": "docx",
        "PDF (.pdf)": "pdf",
        "CSV (.csv)": "csv"
    }
    
    extension = format_mapping.get(file_format, "").lower()
    if not extension:
        st.error("Unsupported file format.")
        return
    
    base_name = os.path.splitext(original_file_name)[0]
    current_date = datetime.datetime.now().strftime("%Y%m%d")
    file_name = f"{base_name}_{current_date}.{extension}"
    
    content_to_save = re.sub(r'<[^>]+>', '', content)
    
    if extension == "pdf":
        save_as_pdf(content_to_save, file_name)
    elif extension == "docx":
        save_as_docx(content_to_save, file_name)
    elif extension == "txt":
        save_as_txt(content_to_save, file_name)
    elif extension == "csv":
        save_as_csv(content_to_save, file_name)

    with open(file_name, "rb") as file:
        st.download_button(
            label=f"Download Document ({extension.upper()})",
            data=file,
            file_name=file_name,
            mime='application/octet-stream'
        )

# Função para baixar o arquivo JSON atualizado
def download_json(data, original_file_name):
    base_name = os.path.splitext(original_file_name)[0]
    current_date = datetime.datetime.now().strftime("%Y%m%d")
    file_name = f"{base_name}_updated_{current_date}.json"
    
    json_data = json.dumps(data, indent=4)
    
    st.download_button(
        label="Download JSON",
        data=json_data,
        file_name=file_name,
        mime="application/json"
    )

# Função principal da aplicação Streamlit
def main():
    st.title("Data Extraction from Documents")
    
    st.sidebar.title("Upload File")
    uploaded_file = st.sidebar.file_uploader("Upload a .docx, .pdf, .txt or .xml file", type=["docx", "pdf", "txt", "xml"])
    
    if uploaded_file:
        file_type = uploaded_file.name.split('.')[-1]
        content = ""
        
        if file_type == 'pdf':
            content = extract_text_from_pdf(uploaded_file)
        elif file_type == 'docx':
            content = extract_text_from_docx(uploaded_file)
        elif file_type == 'txt':
            content = extract_text_from_txt(uploaded_file)
        elif file_type == 'xml':
            content = extract_text_from_xml(uploaded_file)
        else:
            st.error("Unsupported file format.")
            return

        if not content:
            st.error("The file contents are empty.")
            return
        
        variables = parse_key_value_pairs(content)
        
        st.sidebar.subheader("Replace Variables")
        replacements = {}
        for key, value in variables.items():
            replacements[key] = st.sidebar.text_input(key, value=value)
        
        if st.sidebar.button("Apply Substitutions"):
            updated_content = replace_variables(content, replacements)
            st.session_state['content'] = updated_content
            st.session_state['replacements'] = replacements
            
        content_to_display = st.session_state.get('content', content)
        st.write("Updated content:")
        st.markdown(content_to_display, unsafe_allow_html=True)

        # Exibe o JSON atualizado abaixo do conteúdo modificado
        st.write("Updated JSON:")
        json_data = st.session_state.get('replacements', {})
        st.json(json_data)
        
        # Botão para baixar o JSON atualizado
        download_json(json_data, uploaded_file.name)
        
        # Altera o nome do botão para "Download File" de acordo com o formato selecionado
        file_format = st.sidebar.selectbox("Choose a file format to download", ["Texto (.txt)", "DOCX (.docx)", "PDF (.pdf)", "CSV (.csv)"])
        if st.sidebar.button("Download File"):
            download_file(content_to_display, uploaded_file.name, file_format)
            
if __name__ == "__main__":
    main()
